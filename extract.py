import PyPDF2
import glob
import xlrd
from openpyxl import load_workbook
import substring
import csv
import pandas as pd
import pdfkit as pdf
import numpy as np
import string
import re

from docx import Document
import io
import shutil
import os

from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage


from nltk.tokenize import word_tokenize
from nltk.stem.snowball import SnowballStemmer
from nltk.tag import pos_tag
# from nltk.tokenize.moses import MosesDetokenizer
from mosestokenizer import MosesTokenizer, MosesDetokenizer


def xls_to_csv(filename):
  wb  = xlrd.open_workbook("./itineraryData/" + str(filename), on_demand=True )
  for i in range(0, workbook.nsheets-1):
    sheet = wb.sheet_by_index(i)
    print(sheet.name)

    with open("./data/%s.csv" %(sheet.name.replace(" ","") + filename[:-4]), "w+") as file:
        writer = csv.writer(file, delimiter = ",")
        print(sheet, sheet.name, sheet.ncols, sheet.nrows)
        header = [cell.value for cell in sheet.row(0)]
        writer.writerow(header)
        for row_idx in range(1, sheet.nrows):
            row = [int(cell.value) if isinstance(cell.value, float) else cell.value
                   for cell in sheet.row(row_idx)]
            writer.writerow(row)

def csv_to_pdf(filename):
  csv_file = './data/' + filename
  html_file = csv_file[:-3] + 'html'
  pdf_file = csv_file[:-3] + 'pdf'

  df = pd.read_csv(csv_file, sep=',')
  df1 = df.replace(np.nan, '', regex=True)
  df1.to_html(html_file)
  pdf.from_file(html_file, pdf_file)

def csv_to_txt(filename):
  csv_file = './data/' + filename
  # html_file = csv_file[:-3] + 'html'
  # pdf_file = csv_file[:-3] + 'pdf'

  df = pd.read_csv(csv_file, sep=',')
  df1 = df.replace(np.nan, '', regex=True)
  # df1.to_html(html_file)
  # pdf.from_file(html_file, pdf_file)
  tfile = open("./data/%s.txt" %(filename[:-4]), "w+")
  tfile.write(df1.to_string())
  tfile.close()

def convert_pdf_to_txt(filename):
  rsrcmgr = PDFResourceManager()
  retstr = io.StringIO()
  codec = 'utf-8'
  laparams = LAParams()
  device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
  fp = open('./itineraryData/' + filename, 'rb')
  interpreter = PDFPageInterpreter(rsrcmgr, device)
  password = ""
  maxpages = 0
  caching = True
  pagenos = set()

  for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages,
                                password=password,
                                caching=caching,
                                check_extractable=True):
  
    interpreter.process_page(page)

  text = retstr.getvalue()
  print(text)

  fp.close()
  device.close()
  retstr.close()

  text_file = open("./data/%s.txt" %(filename[:-4]), "w+")
  text = re.sub("\s\s+", " ", text)
  text_file.write("%s" % text)
  text_file.close()

def docx_to_txt(filename):
  document = Document("./itineraryData/" + filename)
  textFilename = "./data/" + filename[:-4] + ".txt"
  dates = []
  description = []
  with io.open(textFilename,"w", encoding="utf-8") as textFile:
    for para in document.paragraphs: 
      textFile.write(str(para.text))
    tables = document.tables
    for table in tables:
        for row in table.rows:
            # for cell in row.cells:
          if row.cells[0]:
            for paragraph in row.cells[0].paragraphs:
              if(re.match(r"[a-zA-Z]{4,9} \d+", str(paragraph.text))):
                dates.append(str(paragraph.text))
                textFile.write(str(paragraph.text) + " ")
              else:
                textFile.write(str(paragraph.text) + " ")
          if row.cells[1]:
            print("WHOOOOOOOOOOOOOOOOOO")
            print(row.cells[1].paragraphs[x].text for x in row.cells[1].paragraphs)
            for paragraph in row.cells[1].paragraphs:
              description.append(str(paragraph.text))
              textFile.write(str(paragraph.text) + " ")
  print("BEOOTCH")
  print(dates)
  print(description)
  print("\n\n\n")
  


def clean_txt(filename):
  file = open(filename, 'rt')
  text = file.read()
  file.close()
  # split into words by white space
  tokens = word_tokenize(text)

  with MosesDetokenizer('en') as detokenize:
    detokenize(tokens)

  text_string = " ".join(str(x) for x in tokens)
  # detokenizer = MosesDetokenizer(lang='en')

  # detokenizer.detokenize(tokens)


  # stemming of words
  # snowball = SnowballStemmer("english", ignore_stopwords=True)
  # stemmed = [snowball.stem(word) for word in tokens]
  # print('oink oink im a pig !!!!!!!!!!!!!!!!!!!!!!')
  # print(tokens[:100])
  print(type(text_string))
  print(text_string)
  # dates = re.findall(r'[A-Z][a-z]{1,8}\s\d{1,3}([a-z]{1,3})?,\s\d{2,4}', text_string)
  dates = re.findall(r"[a-zA-Z]{4,9} \d+", text_string)

  print(dates)

pdfCount = len(glob.glob1('./itineraryData',"*.pdf"))
xlsCount = len(glob.glob1('./itineraryData',"*.xls"))

for f in glob.glob1('./itineraryData',"*"):
  s = substring.substringByChar(f, startChar=".")
  file_num = substring.substringByChar(f, endChar='.')
#   if(s == '.xls'):
#     # print("FUCK EXCEL")
#     workbook  = xlrd.open_workbook("./itineraryData/" + str(f), on_demand=True )
#     sheet_number = len(workbook.sheet_names())
#     print("NUMBER OF SHEETS ", sheet_number)
#     xls_to_csv(f)

#     for filename in glob.glob1('./data', "*.csv"):
#       # csv_to_pdf(filename)
#       csv_to_txt(filename)

#     # for filename in glob.glob1('./data', "*.pdf"):
#     #   convert_pdf_to_txt(filename)

#   if(s == '.pdf'):
#     file = open('./itineraryData/' + str(f), 'rb')
#     fileReader = PyPDF2.PdfFileReader(file)
#     print("NUMBER OF PAGES ", fileReader.numPages)
#     # for filename in glob.glob1('./itinerary', "*.pdf"):
#     convert_pdf_to_txt(f)

  if(s == '.docx'):
    docx_to_txt(f)
    clean_txt("./data/" + str(f[:-4]) + '.txt')
    print("hi")


# for f in glob.glob1('./data',"*.txt"):
#   file_num = substring.substringByChar(f, endChar='.')
#   if str(f) == '1.txt':
#     print("hell0")
#     clean_txt("./data/" + str(f))
#     print("bitch")
# for f in range(1, pdfCount):
#   file = open('./itineraryData/' + str(f) + '.pdf', 'rb')

#   # creating a pdf reader object
#   fileReader = PyPDF2.PdfFileReader(file)

#   # print the number of pages in pdf file
#   print(fileReader.numPages)

# for f in range()

# print(xlsCount)
